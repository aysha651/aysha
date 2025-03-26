import streamlit as st
import speech_recognition as sr
from docx import Document
import os
import base64


def add_bg_from_local(image_file):
    with open(image_file, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read())
    st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url(data:image/{"jpg"};base64,{encoded_string.decode()});
        background-size: cover
    }}
    </style>
    """,
    unsafe_allow_html=True
    )
add_bg_from_local('img.jpg') 


def recognize_speech(prompt):
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.write(prompt)
        audio = r.listen(source)

        try:
            text = r.recognize_google(audio)
            return text.lower().strip()
        except sr.UnknownValueError:
            st.error("Sorry, I could not understand the audio.")
            return None
        except sr.RequestError:
            st.error("Could not request results from Google Speech Recognition service.")
            return None

def create_document():
    return Document()

def add_paragraph(doc, text):
    doc.add_paragraph(text)
    st.write(f"Added paragraph: {text}")  # Debugging step

def save_document(doc, filename):
    doc.save(filename)
    st.write(f"Document saved as: {filename}")  # Debugging step

def open_document(filename):
    if os.name == 'nt':
        os.system(f'start "" \"{filename}\"')
    elif os.name == 'posix':
        os.system(f'open \"{filename}\"')
    else:
        st.error("Unsupported operating system.")
    st.write(f"Opening document: {filename}")  # Debugging step

st.title("Voice-Controlled Word Document Manager")
st.write("_______________________________________")
st.write("Please say a command (e.g., 'create document', 'add paragraph', 'save document', 'open document')")

# Initialize a session state to keep track of documents
if 'documents' not in st.session_state:
    st.session_state.documents = {}

# Continuous listening loop
while True:
    command = recognize_speech("Listening for commands...")

    if command:
        st.success(f"You said: {command}")

        if "create document" in command:
            doc_name = recognize_speech("Please say the document name:")
            if doc_name:
                st.session_state.documents[doc_name] = create_document()
                st.success(f"Document '{doc_name}' created.")
            else:
                st.error("No document name recognized.")

        elif "add paragraph" in command:
            doc_name = recognize_speech("Please say the document name to add the paragraph to:")
            if doc_name and doc_name in st.session_state.documents:
                content = recognize_speech("Please say the paragraph content:")
                if content:
                    add_paragraph(st.session_state.documents[doc_name], content)
                    st.success("Paragraph added.")
                else:
                    st.error("No content recognized.")
            else:
                st.error("Document not found. Please create the document first.")

        elif "save document" in command:
            doc_name = recognize_speech("Please say the document name to save:")
            if doc_name and doc_name in st.session_state.documents:
                filename = f"{doc_name}.docx"
                save_document(st.session_state.documents[doc_name], filename)
                st.success(f"Document saved as '{filename}'.")
            else:
                st.error("Document not found. Please create the document first.")

        elif "open document" in command:
            doc_name = recognize_speech("Please say the document name to open:")
            filename = f"{doc_name}.docx"
            if os.path.exists(filename):
                open_document(filename)
                st.success(f"Document '{filename}' opened.")
            else:
                st.error("Document not found.")

        elif "exit" in command:
            st.success("Exiting the application.")
            break
